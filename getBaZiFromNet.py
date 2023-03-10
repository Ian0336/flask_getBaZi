# coding:utf-8
from ast import If
from cgitb import small
from html.entities import entitydefs
from itertools import count
from msilib.schema import ComboBox, Font
from operator import truediv
from pickle import TRUE
from sre_parse import State
from xml.dom.minidom import Entity
""" from isort import place_module """
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from bs4 import BeautifulSoup
""" from pykeyboard import PyKeyboard
from pymouse import PyMouse """
import re
""" import requests """
import hashlib
import os
import time
import glob
import tkinter as tk
from tkinter import VERTICAL, ttk
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
import smtplib
from datetime import date
from docx2pdf import convert
import pythoncom


def crawler():
    global bigFate
    global bigFateText
    global loveGod
    global body

    print(isLunar,day,mon)
    option = webdriver.EdgeOptions()
    option.add_argument("headless")  
    driver = webdriver.Edge(executable_path="D:\\edgedriver\\msedgedriver", options=option)
    driver.set_window_size(1552, 893)
    driver.get("https://myfate.herokuapp.com/")
    if isLunar:
        driver.find_element("id", "lunar").click()
    if isMan:
        driver.find_element("id", "male").click()
    select = Select(driver.find_element("id", "Year"))
    select.select_by_index(year - 1900)
    select = Select(driver.find_element("id", "Month"))
    select.select_by_index(mon - 1)
    select = Select(driver.find_element("id", "Day"))
    select.select_by_index(day - 1)
    select = Select(driver.find_element("id", "Hour"))
    select.select_by_index(hr)
    if(isLunar and isLeapMonth):
        driver.find_element(By.NAME, "Leap").click()
    driver.find_element(By.TAG_NAME, "button").click()
    element = WebDriverWait(driver, 10).until(
        EC.presence_of_element_located(
            (By.CLASS_NAME, "card-header"))
    )
    print("OK")
    firstData_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'card my-3 shadow'][last()]//div[@class = 'flexcontainer']//span[@class = 'v']")
    for i in range(0, len(firstData_obj)):
        """ print(i)
        print(str((firstData_obj[i].text))) """
        if i > 0 and i < 5:
            mainStar.append(str((firstData_obj[i].text)))
        if i > 5 and i < 18:
            secondStar[(i - 6) // 3].append(str((firstData_obj[i].text)))
        if i > 18 and i < 23:
            fate.append(str((firstData_obj[i].text)))
    secondData_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'card my-3 shadow'][last()]//div[@class = 'flexcontainer']//span[@class = 'vs']")
    for i in range(0, len(secondData_obj)):
        if i > 0:
            bury[(i - 1) // 3].append(str((secondData_obj[i].text)))
    getLen_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'card my-3 shadow'][last()]//div[@class = 'flexcontainer']//div[@class = 'subcontainer'][last()]/div[@class='s']")
    getlen = []
    for i in getLen_obj:
        getlen.append(len(str(i.text).replace('\n', '')))
    third_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'card my-3 shadow'][last()]//div[@class = 'flexcontainer']//div[@class = 'subcontainer'][last()]//span[@class='vl']")
    index = 0
    sumLen = [0, 0, 0, 0]
    allshansha = []
    for i in range(0, len(third_obj)):
        if sumLen[index] == getlen[index]:
            index += 1
        allshansha.append(str((third_obj[i].text)))
        shansha[index].append(str((third_obj[i].text)))
        sumLen[index] += len((str((third_obj[i].text))))
    eightword_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'card my-3 shadow'][last()]//div[@class = 'flexcontainer']//div[@class = 'pillar']")
    for i in range(0, len(eightword_obj)):
        eightword_f.append(str((eightword_obj[i].text))[0])
        eightword_s.append(str((eightword_obj[i].text))[1])
    othershansha_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'card my-3 shadow'][last()]/div[@class = 'pane']/div[@class = 'mx-3']//td")
    for i in range(0, len(othershansha_obj)):
        if i % 2 == 0 and str((othershansha_obj[i].text)) not in allshansha:
            othershansha.append(str((othershansha_obj[i].text)))
    bigFate_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'card my-3 shadow'][last()-1]//div[@class = 'pane']/div[@class = 'flexcontainer'][1]/div[@class = 'subcontainer'][last()-1]/div[@class = 's'][1]/span[1]")
    bigFate = str(bigFate_obj[0].text)
    bigFateText_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'card my-3 shadow'][last()-1]//div[@class = 'pane']/div[@class = 'flexcontainer'][1]/div[@class = 'subcontainer'][last()]")
    bigFateText = str(bigFateText_obj[0].text)
    bigFateYears_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'card my-3 shadow'][last()-1]//div[@class = 'pane']/div[@class = 'flexcontainer'][1]/div[@class = 'subcontainer'][last()-1]/div[@class = 's']/span[@class = 'v']")
    for i in range(0, len(bigFateYears_obj)):
        if i < 10:
            bigFateYears.append(str(bigFateYears_obj[i].text))
    loveGod_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'container px-4'][3]/div[@class = 'card my-3 shadow'][2]/div[@class = 'card-body']")
    loveGod = str(loveGod_obj[0].text).split('\n')[3]
    body_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'container px-4'][3]/div[@class = 'card my-3 shadow'][5]//div[@class = 'pane']/div[@class = 'flexcontainer']/div[@class = 'hcontainer'][3]")
    body = str(body_obj[0].text)
    personality_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'container px-4'][last()-2]//div[@class = 'card my-3 shadow'][last()]//td")
    for i in range(0, len(personality_obj), 2):
        personality[str(personality_obj[i].text)] = str(
            personality_obj[i+1].text)
    horseFlower_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'container px-4'][last()-1]//table[@class ='table table-striped table-bordered']//td")
    for i in range(0, len(horseFlower_obj), 2):
        horseFlower[str(horseFlower_obj[i].text)] = str(
            horseFlower_obj[i+1].text)
    shanshaExplain_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'container px-4'][last()]//table[@class ='table table-striped table-bordered']//td")
    for i in range(0, len(shanshaExplain_obj), 2):
        shanshaExplain[str(shanshaExplain_obj[i].text)] = str(
            shanshaExplain_obj[i+1].text)

    bornDay_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'container px-4'][last()]//div[@class = 'flexcontainer'][1]//div[@class = 'hcontainer']")
    for i in range(0, len(bornDay_obj)):

        bornDay.append(str((bornDay_obj[i].text).replace('\n', '')))
    driver.quit()





def five_g(i):
    if(i == 1 or i == 0):
        return "???"
    if(i == 3 or i == 2):
        return "???"
    if(i == 5 or i == 4):
        return "???"
    if(i == 7 or i == 6):
        return "???"
    if(i == 9 or i == 8):
        return "???"


def five_z(i):
    if(i == 3 or i == 2):
        return "???"
    if(i == 6 or i == 5):
        return "???"
    if(i == 4 or i == 10 or i == 1 or i == 7):
        return "???"
    if(i == 8 or i == 9):
        return "???"
    if(i == 0 or i == 11):
        return "???"


def adjust():
    global couterAnimal
    global yourAnimal
    global sixCan
    a = ['???', '???', '???', '???', '???', '???', '???', '???', '???', '???']
    b = ['???', '???', '???', '???', '???', '???', '???', '???', '???', '???', '???', '???']
    e = []
    for i in range(60):
        e.append(a[i % 10]+b[i % 12])
    tmp = eightword_f[0]+eightword_s[0]
    for i in range(60):
        if tmp == e[i]:
            for j in range(100):
                flowYears.append(e[(i+j) % 60])
            break
    c = {}
    d = {}
    for i in range(10):
        c[a[i]] = i
    for i in range(12):
        d[b[i]] = i
    g = []
    z = []
    for i in range(4):
        g.append(c[eightword_f[i]])
        z.append(d[eightword_s[i]])
        eightword_f_elem.append(five_g(g[i]))
        eightword_s_elem.append(five_z(z[i]))
    yourAnimal = z[0]
    counterAnimal_list = [100, 8, 7, 6, 5, 4, 3, 2, 1, 12, 11, 10, 9]
    couterAnimal = counterAnimal_list[yourAnimal+1]-1

    sixZi = eightword_s
    todayYear = date.today().year
    age = todayYear - year + 1
    tmp = ""

    for i in range(0, len(bigFateYears)):
        if(age >= (i*10) + int(bigFate)):
            tmp = bigFateYears[i]
        else:
            break

    if(tmp != ""):
        sixZi.append(tmp[0])
        sixZi.append(tmp[1])
    sixZi.append(flowYears[age][0])
    sixZi.append(flowYears[age][1])

    if all(elem in sixZi for elem in ['???', '???', '???']):
        sixCan += "'???','???','???'???????????? /"
    if all(elem in sixZi for elem in ['???', '???', '???']):
        sixCan += "'???','???','???'???????????? /"
    if all(elem in sixZi for elem in ['???', '???', '???']):
        sixCan += "'???','???','???''???????????? /"
    if all(elem in sixZi for elem in ['???', '???', '???']):
        sixCan += "'???','???','???'???????????? /"
    sixZi = sixZi + eightword_f
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? / "

    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? / "

    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "??????????????? / "

    if all(elem in sixZi for elem in ['???', '???', '???']):
        sixCan += "????????????????????? / "
    if all(elem in sixZi for elem in ['???', '???', '???']):
        sixCan += "????????????????????? / "
    if all(elem in sixZi for elem in ['???', '???', '???']):
        sixCan += "????????????????????? / "
    if all(elem in sixZi for elem in ['???', '???', '???']):
        sixCan += "????????????????????? / "
    if all(elem in sixZi for elem in ['???', '???', '???', '???']):
        sixCan += "????????????????????? / "

    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? / "

    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "????????? / "
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "????????? /????????? /"
    cnt = 0
    for i in sixZi:
        if i == '???':
            cnt += 1
    if cnt == 2:
        sixCan += "????????? /"
    cnt = 0
    for i in sixZi:
        if i == '???':
            cnt += 1
    if cnt == 2:
        sixCan += "????????? /"
    cnt = 0
    for i in sixZi:
        if i == '???':
            cnt += 1
    if cnt == 2:
        sixCan += "????????? /"
    cnt = 0
    for i in sixZi:
        if i == '???':
            cnt += 1
    if cnt == 2:
        sixCan += "????????? /"

    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? /"
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? /"
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? /"
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? /"
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? /"
    if all(elem in sixZi for elem in ['???', '???']):
        sixCan += "???????????? /"

    listy = []
    listm = []
    listd = []
    listh = []
    if z[1] == 0:
        if z[2] == 5:
            listd.append("??????")
        elif g[2] == 8:
            listd.append("??????")
    if z[1] == 1:
        if g[2] == 6:
            listd.append("??????")
        elif g[2] == 6:
            listd.append("??????")
    if z[1] == 2:
        if g[2] == 3:
            listd.append("??????")
        elif g[2] == 2:
            listd.append("??????")
    if z[1] == 3:
        if z[2] == 8:
            listd.append("??????")
        elif g[2] == 0:
            listd.append("??????")
    if z[1] == 4:
        if g[2] == 8:
            listd.append("??????")
        elif g[2] == 8:
            listd.append("??????")
    if z[1] == 5:
        if g[2] == 7:
            listd.append("??????")
        elif g[2] == 6:
            listd.append("??????")
    if z[1] == 6:
        if z[2] == 11:
            listd.append("??????")
        elif g[2] == 2:
            listd.append("??????")
    if z[1] == 7:
        if g[2] == 0:
            listd.append("??????")
        elif g[2] == 0:
            listd.append("??????")
    if z[1] == 8:
        if g[2] == 9:
            listd.append("??????")
        elif g[2] == 8:
            listd.append("??????")
    if z[1] == 9:
        if z[2] == 2:
            listd.append("??????")
        elif g[2] == 6:
            listd.append("??????")
    if z[1] == 10:
        if g[2] == 2:
            listd.append("??????")
        elif g[2] == 2:
            listd.append("??????")
    if z[1] == 11:
        if g[2] == 1:
            listd.append("??????")
        elif g[2] == 0:
            listd.append("??????")

    if z[0] == 0 or z[0] == 1 or z[0] == 11:
        for i in range(4):
            if z[i] == 2:
                if i == 0:
                    listy.append("??????")
                elif i == 1:
                    listm.append("??????")
                elif i == 2:
                    listd.append("??????")
                else:
                    listh.append("??????")
            if z[i] == 10:
                if i == 0:
                    listy.append("??????")
                elif i == 1:
                    listm.append("??????")
                elif i == 2:
                    listd.append("??????")
                else:
                    listh.append("??????")
    if z[0] == 2 or z[0] == 3 or z[0] == 4:
        for i in range(4):
            if z[i] == 5:
                if i == 0:
                    listy.append("??????")
                elif i == 1:
                    listm.append("??????")
                elif i == 2:
                    listd.append("??????")
                else:
                    listh.append("??????")
            if z[i] == 1:
                if i == 0:
                    listy.append("??????")
                elif i == 1:
                    listm.append("??????")
                elif i == 2:
                    listd.append("??????")
                else:
                    listh.append("??????")
    if z[0] == 5 or z[0] == 6 or z[0] == 7:
        for i in range(4):
            if z[i] == 8:
                if i == 0:
                    listy.append("??????")
                elif i == 1:
                    listm.append("??????")
                elif i == 2:
                    listd.append("??????")
                else:
                    listh.append("??????")
            if z[i] == 4:
                if i == 0:
                    listy.append("??????")
                elif i == 1:
                    listm.append("??????")
                elif i == 2:
                    listd.append("??????")
                else:
                    listh.append("??????")
    if z[0] == 8 or z[0] == 9 or z[0] == 10:
        for i in range(4):
            if z[i] == 11:
                if i == 0:
                    listy.append("??????")
                elif i == 1:
                    listm.append("??????")
                elif i == 2:
                    listd.append("??????")
                else:
                    listh.append("??????")
            if z[i] == 7:
                if i == 0:
                    listy.append("??????")
                elif i == 1:
                    listm.append("??????")
                elif i == 2:
                    listd.append("??????")
                else:
                    listh.append("??????")

    if z[0] == 0 or z[0] == 4 or z[0] == 8:
        if z[1] == 2:
            listm.append("????????????")

        if z[1] == 1:
            listm.append("????????????")
    if z[0] == 1 or z[0] == 5 or z[0] == 9:
        if z[1] == 7:
            listm.append("????????????")
        if z[1] == 10:
            listm.append("????????????")
    if z[0] == 2 or z[0] == 6 or z[0] == 10:
        if z[1] == 5:
            listm.append("????????????")
        if z[1] == 8:
            listm.append("????????????")
    if z[0] == 3 or z[0] == 7 or z[0] == 11:
        if z[1] == 3:
            listm.append("????????????")
        if z[1] == 9:
            listm.append("????????????")

    if 0 in g:
        if 4 in g:
            if 6 in g:
                listy.append("????????????")
    if 1 in g:
        if 2 in g:
            if 3 in g:
                listy.append("????????????")
    if 7 in g:
        if 8 in g:
            if 9 in g:
                listy.append("????????????")

    if g[2] == 6 and z[2] == 4:
        listy.append("??????")
    if g[2] == 8 and z[2] == 4:
        listy.append("??????")
    if g[2] == 6 and z[2] == 10:
        listy.append("??????")
    if g[2] == 4 and z[2] == 10:
        listy.append("??????")

    if g[2] == 1 and z[2] == 5:
        listy.append("?????????")
    if g[2] == 3 and z[2] == 5:
        listy.append("?????????")
    if g[2] == 0 and z[2] == 2:
        listy.append("?????????")
    if g[2] == 8 and z[2] == 2:
        listy.append("?????????")
    if g[2] == 2 and z[2] == 6:
        listy.append("?????????")
    if g[2] == 4 and z[2] == 6:
        listy.append("?????????")
    if g[2] == 8 and z[2] == 0:
        listy.append("?????????")
    if g[2] == 7 and z[2] == 11:
        listy.append("?????????")
    if g[2] == 4 and z[2] == 8:
        listy.append("?????????")

    def check():
        g = listg
        z = listz

        if g[2] == 0:
            for i in range(4):
                if z[i] == 1 or z[i] == 7:
                    if i == 0:
                        listy.append("????????????")
                    elif i == 1:
                        listm.append("????????????")
                    elif i == 2:
                        listd.append("????????????")
                    else:
                        listh.append("????????????")
                if z[i] == 5:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 2:
                    if i == 0:
                        listy.append("?????????")
                    elif i == 1:
                        listm.append("?????????")
                    elif i == 2:
                        listd.append("?????????")
                    else:
                        listh.append("?????????")
                if z[i] == 3:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 6:
                    if i == 0:
                        listy.append("?????????")
                    elif i == 1:
                        listm.append("?????????")
                    elif i == 2:
                        listd.append("?????????")
                    else:
                        listh.append("?????????")
                if z[i] == 9:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
        if g[2] == 1:
            for i in range(4):
                if z[i] == 0 or z[i] == 8:
                    if i == 0:
                        listy.append("????????????")
                    elif i == 1:
                        listm.append("????????????")
                    elif i == 2:
                        listd.append("????????????")
                    else:
                        listh.append("????????????")
                if z[i] == 6:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 3:
                    if i == 0:
                        listy.append("?????????")
                    elif i == 1:
                        listm.append("?????????")
                    elif i == 2:
                        listd.append("?????????")
                    else:
                        listh.append("?????????")
                if z[i] == 4:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 6:
                    if i == 0:
                        listy.append("?????????")
                    elif i == 1:
                        listm.append("?????????")
                    elif i == 2:
                        listd.append("?????????")
                    else:
                        listh.append("?????????")
                if z[i] == 10:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
        if g[2] == 2:
            for i in range(4):
                if z[i] == 11 or z[i] == 9:
                    if i == 0:
                        listy.append("????????????")
                    elif i == 1:
                        listm.append("????????????")
                    elif i == 2:
                        listd.append("????????????")
                    else:
                        listh.append("????????????")
                if z[i] == 8:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 5:
                    if i == 0:
                        listy.append("?????????")
                    elif i == 1:
                        listm.append("?????????")
                    elif i == 2:
                        listd.append("?????????")
                    else:
                        listh.append("?????????")
                if z[i] == 6:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 2:
                    if i == 0:
                        listy.append("?????????")
                    elif i == 1:
                        listm.append("?????????")
                    elif i == 2:
                        listd.append("?????????")
                    else:
                        listh.append("?????????")
                if z[i] == 0:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
        if g[2] == 3:
            for i in range(4):
                if z[i] == 11 or z[i] == 9:
                    if i == 0:
                        listy.append("????????????")
                    elif i == 1:
                        listm.append("????????????")
                    elif i == 2:
                        listd.append("????????????")
                    else:
                        listh.append("????????????")
                if z[i] == 9:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 6:
                    if i == 0:
                        listy.append("?????????")
                    elif i == 1:
                        listm.append("?????????")
                    elif i == 2:
                        listd.append("?????????")
                    else:
                        listh.append("?????????")
                if z[i] == 7:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 7:
                    if i == 0:
                        listy.append("?????????")
                    elif i == 1:
                        listm.append("?????????")
                    elif i == 2:
                        listd.append("?????????")
                    else:
                        listh.append("?????????")
                if z[i] == 1:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
        if g[2] == 4:
            for i in range(4):
                if z[i] == 7 or z[i] == 1:
                    if i == 0:
                        listy.append("????????????")
                    elif i == 1:
                        listm.append("????????????")
                    elif i == 2:
                        listd.append("????????????")
                    else:
                        listh.append("????????????")
                if z[i] == 8:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 5:
                    if i == 0:
                        listy.append("?????????")
                    elif i == 1:
                        listm.append("?????????")
                    elif i == 2:
                        listd.append("?????????")
                    else:
                        listh.append("?????????")
                if z[i] == 6:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 4:
                    if i == 0:
                        listy.append("?????????")
                    elif i == 1:
                        listm.append("?????????")
                    elif i == 2:
                        listd.append("?????????")
                    else:
                        listh.append("?????????")
                if z[i] == 0:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
        if g[2] == 5:
            for i in range(4):
                if z[i] == 0 or z[i] == 8:
                    if i == 0:
                        listy.append("????????????")
                    elif i == 1:
                        listm.append("????????????")
                    elif i == 2:
                        listd.append("????????????")
                    else:
                        listh.append("????????????")
                if z[i] == 9:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 6:
                    if i == 0:
                        listy.append("?????????")
                    elif i == 1:
                        listm.append("?????????")
                    elif i == 2:
                        listd.append("?????????")
                    else:
                        listh.append("?????????")
                if z[i] == 7:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 4:
                    if i == 0:
                        listy.append("?????????")
                    elif i == 1:
                        listm.append("?????????")
                    elif i == 2:
                        listd.append("?????????")
                    else:
                        listh.append("?????????")
                if z[i] == 1:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
        if g[2] == 6:
            for i in range(4):
                if z[i] == 1 or z[i] == 7:
                    if i == 0:
                        listy.append("????????????")
                    elif i == 1:
                        listm.append("????????????")
                    elif i == 2:
                        listd.append("????????????")
                    else:
                        listh.append("????????????")
                if z[i] == 11:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 8:
                    if i == 0:
                        listy.append("?????????")
                    elif i == 1:
                        listm.append("?????????")
                    elif i == 2:
                        listd.append("?????????")
                    else:
                        listh.append("?????????")
                if z[i] == 9:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 10:
                    if i == 0:
                        listy.append("?????????")
                    elif i == 1:
                        listm.append("?????????")
                    elif i == 2:
                        listd.append("?????????")
                    else:
                        listh.append("?????????")
                if z[i] == 3:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
        if g[2] == 7:
            for i in range(4):
                if z[i] == 2 or z[i] == 6:
                    if i == 0:
                        listy.append("????????????")
                    elif i == 1:
                        listm.append("????????????")
                    elif i == 2:
                        listd.append("????????????")
                    else:
                        listh.append("????????????")
                if z[i] == 0:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 9:
                    if i == 0:
                        listy.append("?????????")
                    elif i == 1:
                        listm.append("?????????")
                    elif i == 2:
                        listd.append("?????????")
                    else:
                        listh.append("?????????")
                if z[i] == 10:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 9:
                    if i == 0:
                        listy.append("?????????")
                    elif i == 1:
                        listm.append("?????????")
                    elif i == 2:
                        listd.append("?????????")
                    else:
                        listh.append("?????????")
                if z[i] == 4:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
        if g[2] == 8:
            for i in range(4):
                if z[i] == 3 or z[i] == 5:
                    if i == 0:
                        listy.append("????????????")
                    elif i == 1:
                        listm.append("????????????")
                    elif i == 2:
                        listd.append("????????????")
                    else:
                        listh.append("????????????")
                if z[i] == 2:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 11:
                    if i == 0:
                        listy.append("?????????")
                    elif i == 1:
                        listm.append("?????????")
                    elif i == 2:
                        listd.append("?????????")
                    else:
                        listh.append("?????????")
                if z[i] == 0:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 0:
                    if i == 0:
                        listy.append("?????????")
                    elif i == 1:
                        listm.append("?????????")
                    elif i == 2:
                        listd.append("?????????")
                    else:
                        listh.append("?????????")
                if z[i] == 6:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
        if g[2] == 9:
            for i in range(4):
                if z[i] == 5 or z[i] == 3:
                    if i == 0:
                        listy.append("????????????")
                    elif i == 1:
                        listm.append("????????????")
                    elif i == 2:
                        listd.append("????????????")
                    else:
                        listh.append("????????????")
                if z[i] == 3:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 0:
                    if i == 0:
                        listy.append("?????????")
                    elif i == 1:
                        listm.append("?????????")
                    elif i == 2:
                        listd.append("?????????")
                    else:
                        listh.append("?????????")
                if z[i] == 1:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 8:
                    if i == 0:
                        listy.append("?????????")
                    elif i == 1:
                        listm.append("?????????")
                    elif i == 2:
                        listd.append("?????????")
                    else:
                        listh.append("?????????")
                if z[i] == 7:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")

        if z[2] == 0:
            for i in range(4):
                if z[i] == 0:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 4:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 2:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 5:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 11:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 9:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
        if z[2] == 1:
            for i in range(4):
                if z[i] == 9:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 1:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 11:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 2:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 8:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 6:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
        if z[2] == 2:
            for i in range(4):
                if z[i] == 6:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 1:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 8:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 11:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 5:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 3:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
        if z[2] == 3:
            for i in range(4):
                if z[i] == 3:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 7:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 5:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 8:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 2:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 0:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
        if z[2] == 4:
            for i in range(4):
                if z[i] == 0:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 4:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 2:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 5:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 11:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 9:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
        if z[2] == 5:
            for i in range(4):
                if z[i] == 9:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 1:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 11:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 2:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 8:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 6:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
        if z[2] == 6:
            for i in range(4):
                if z[i] == 6:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 10:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 8:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 11:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 5:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 3:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
        if z[2] == 7:
            for i in range(4):
                if z[i] == 3:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 7:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 5:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 8:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 2:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 0:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
        if z[2] == 8:
            for i in range(4):
                if z[i] == 0:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 4:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 2:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 5:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 11:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 9:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
        if z[2] == 9:
            for i in range(4):
                if z[i] == 9:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 1:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 11:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 2:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 8:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 6:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
        if z[2] == 10:
            for i in range(4):
                if z[i] == 6:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 10:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 8:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 11:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 5:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 3:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
        if z[2] == 11:
            for i in range(4):
                if z[i] == 3:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 7:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 5:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 8:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 2:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")
                if z[i] == 0:
                    if i == 0:
                        listy.append("??????")
                    elif i == 1:
                        listm.append("??????")
                    elif i == 2:
                        listd.append("??????")
                    else:
                        listh.append("??????")

    shansha[0] = shansha[0] + listy
    shansha[1] = shansha[1] + listm
    shansha[2] = shansha[2] + listd
    shansha[3] = shansha[3] + listh
    listy.clear()
    listm.clear()
    listd.clear()
    listh.clear()
    # ????????????

    for i in range(4):
        if z[i] == 5:
            if g[2] == 0 or g[2] == 5:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 0:
            if g[2] == 1 or g[2] == 6:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 8:
            if g[2] == 2 or g[2] == 7:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 10:
            if g[2] == 3 or g[2] == 8:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 3:
            if g[2] == 4 or g[2] == 9:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 1 or z[i] == 9 or z[i] == 5:
            if g[2] == 0 or g[2] == 5:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 0:
            if g[2] == 1 or g[2] == 2 or g[2] == 3 or g[2] == 4:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 7 or z[i] == 9 or z[i] == 11:
            if g[2] == 6:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 10 or z[i] == 6 or z[i] == 2:
            if g[2] == 7 or g[2] == 8 or g[2] == 9:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 4 or z[i] == 0 or z[i] == 8:
            if g[2] == 0 or g[2] == 1 or g[2] == 2 or g[2] == 3:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 7 or z[i] == 3 or z[i] == 11:
            if g[2] == 4 or g[2] == 5 or g[2] == 6:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 10 or z[i] == 6 or z[i] == 2:
            if g[2] == 7 or g[2] == 8 or g[2] == 9:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 1 or z[i] == 6:
            if g[2] == 0 or g[2] == 1:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 0:
            if g[2] == 2 or g[2] == 3 or g[2] == 8 or g[2] == 9:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 10:
            if g[2] == 4:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 6:
            if g[2] == 5:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 2:
            if g[2] == 6 or g[2] == 7:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 7 or z[i] == 8:
            if g[2] == 0 or g[2] == 1:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 4 or z[i] == 5:
            if g[2] == 2 or g[2] == 3:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 2:
            if g[2] == 6 or g[2] == 7:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 1:
            if g[2] == 8 or g[2] == 9:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 6 or z[i] == 7:
            if g[2] == 0 or g[2] == 1:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 8 or z[i] == 9:
            if g[2] == 2 or g[2] == 3:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 10 or z[i] == 5:
            if g[2] == 4 or g[2] == 5:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 2:
            if g[2] == 6 or g[2] == 7:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 11 or z[i] == 1:
            if g[2] == 8 or g[2] == 9:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 8 or z[i] == 9:
            if g[2] == 0 or g[2] == 1:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 0 or z[i] == 11:
            if g[2] == 2 or g[2] == 3:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 3 or z[i] == 2:
            if g[2] == 4 or g[2] == 5:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 6 or z[i] == 5:
            if g[2] == 6 or g[2] == 7:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 7 or z[i] == 1 or z[i] == 10 or z[i] == 4:
            if g[2] == 8 or g[2] == 9:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 8 or z[i] == 7:
            if g[2] == 2 or g[2] == 3:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 4:
            if g[2] == 0 or g[2] == 1:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 2:
            if g[2] == 4 or g[2] == 5:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 10:
            if g[2] == 6 or g[2] == 7:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 1:
            if g[2] == 8 or g[2] == 9:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 9:
            if g[2] == 0 or g[2] == 1:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 0:
            if g[2] == 2 or g[2] == 3:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 6:
            if g[2] == 4 or g[2] == 5 or g[2] == 8 or g[2] == 9:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 3:
            if g[2] == 6 or g[2] == 7:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    shansha[3] = shansha[3] + listh
    listh.clear()

    for i in range(4):
        if z[i] == 9:
            if z[2] == 0:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 6:
            if z[2] == 1:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 7:
            if z[2] == 2:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 8:
            if z[2] == 3:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 11:
            if z[2] == 4:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 10:
            if z[2] == 5:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 1:
            if z[2] == 6:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 2:
            if z[2] == 7:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 3:
            if z[2] == 8:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 0:
            if z[2] == 9:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 5:
            if z[2] == 10:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 4:
            if z[2] == 11:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    if z[3] == (z[2]+4) % 12:
        listh.append("?????????")
    if z[2] == (z[3]+2) % 12:
        listh.append("?????????")

    for i in range(4):
        if z[i] == 1:
            if z[2] == 0 or z[2] == 3 or z[2] == 6 or z[2] == 9:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 3:
            if z[2] == 1 or z[2] == 4 or z[2] == 7 or z[2] == 10:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 8:
            if z[2] == 2 or z[2] == 5 or z[2] == 8 or z[2] == 11:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 5 or z[i] == 6:
            if z[2] == 0 or z[2] == 4 or z[2] == 8:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 0 or z[i] == 3:
            if z[2] == 1 or z[2] == 5 or z[2] == 9:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 6 or z[i] == 4:
            if z[2] == 2 or z[2] == 6 or z[2] == 10:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 8 or z[i] == 6:
            if z[2] == 3 or z[2] == 7 or z[2] == 11:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 5:
            if z[2] == 0 or z[2] == 4 or z[2] == 8:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 2:
            if z[2] == 1 or z[2] == 5 or z[2] == 9:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 4:
            if z[2] == 2 or z[2] == 6 or z[2] == 10:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 7:
            if z[2] == 3 or z[2] == 7 or z[2] == 11:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 6 or z[i] == 1 or z[i] == 10 or z[i] == 4:
            if z[2] == 0 or z[2] == 3 or z[2] == 6 or z[2] == 9:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 9 or z[i] == 3 or z[i] == 6 or z[i] == 0:
            if z[2] == 1 or z[2] == 4 or z[2] == 7 or z[2] == 10:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 11 or z[i] == 5 or z[i] == 8 or z[i] == 2:
            if z[2] == 2 or z[2] == 5 or z[2] == 8 or z[2] == 11:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 6:
            if z[2] == 0 or z[2] == 3 or z[2] == 6 or z[2] == 9:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 6:
            if z[2] == 1 or z[2] == 4 or z[2] == 7 or z[2] == 10:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 2:
            if z[2] == 2 or z[2] == 5 or z[2] == 8 or z[2] == 11:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 2:
            if z[2] == 0 or z[2] == 3 or z[2] == 6 or z[2] == 9:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 7:
            if z[2] == 1 or z[2] == 4 or z[2] == 7 or z[2] == 10:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 9:
            if z[2] == 2 or z[2] == 5 or z[2] == 8 or z[2] == 11:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 5:
            if z[2] == 0 or z[2] == 2:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 6:
            if z[2] == 1 or z[2] == 10:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 0:
            if z[2] == 3:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 6:
            if z[2] == 4 or z[2] == 5 or z[2] == 8:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 7:
            if z[2] == 6 or z[2] == 7:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 11:
            if z[2] == 9 or z[2] == 11:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
    shansha[3] = shansha[3] + listh
    listh.clear()
    listd.clear()
    listy.clear()
    listm.clear()

    for i in range(4):
        if z[i] == 6:
            if z[1] == 2 or z[1] == 3:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 7:
            if z[1] == 4 or z[1] == 5:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 10 or z[i] == 3:
            if z[1] == 6 or z[1] == 7:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 8 or z[i] == 5:
            if z[1] == 8 or z[1] == 9:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 3 or z[i] == 2:
            if z[1] == 10 or z[1] == 11:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 9 or z[i] == 4:
            if z[1] == 0 or z[1] == 1:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 10 or z[i] == 7:
            if z[1] == 2 or z[1] == 3 or z[1] == 4:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 4 or z[i] == 1:
            if z[1] == 6 or z[1] == 5 or z[1] == 7 or z[1] == 8 or z[1] == 9 or z[1] == 10:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 4 or z[i] == 7:
            if z[1] == 11 or z[1] == 0 or z[1] == 1:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 8 or z[i] == 2:
            if z[1] == 2 or z[1] == 3 or z[1] == 4:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 7:
            if z[1] == 5 or z[1] == 6 or z[1] == 7:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 9:
            if z[1] == 8 or z[1] == 9 or z[1] == 10:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 1:
            if z[1] == 11 or z[1] == 0 or z[1] == 1:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 11 or z[i] == 5:
            if z[1] == 2 or z[1] == 8:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 10 or z[i] == 4:
            if z[1] == 3 or z[1] == 9:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 9 or z[i] == 3:
            if z[1] == 4 or z[1] == 10:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 8 or z[i] == 2:
            if z[1] == 5 or z[1] == 11:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 7 or z[i] == 1:
            if z[1] == 6 or z[1] == 0:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 6 or z[i] == 0:
            if z[1] == 7 or z[1] == 1:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 4 or z[i] == 10 or z[i] == 9:
            if z[1] == 2 or z[1] == 3 or z[1] == 4:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 0 or z[i] == 3 or z[i] == 7:
            if z[1] == 5 or z[1] == 6 or z[1] == 7:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 1 or z[i] == 2 or z[i] == 6:
            if z[1] == 8 or z[1] == 9 or z[1] == 10:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 5 or z[i] == 8 or z[i] == 11:
            if z[1] == 11 or z[1] == 0 or z[1] == 1:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 4:
            if z[1] == 2 or z[1] == 3 or z[1] == 4:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 7:
            if z[1] == 5 or z[1] == 6 or z[1] == 7:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 10:
            if z[1] == 8 or z[1] == 9 or z[1] == 10:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 1:
            if z[1] == 11 or z[1] == 0 or z[1] == 1:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 2:
            if z[1] == 2:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 3:
            if z[1] == 3:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 8:
            if z[1] == 4:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 1:
            if z[1] == 5:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 10:
            if z[1] == 6:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 9:
            if z[1] == 7:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 4:
            if z[1] == 8:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 5:
            if z[1] == 9:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 6:
            if z[1] == 10:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 7:
            if z[1] == 11:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 11:
            if z[1] == 0:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 0:
            if z[1] == 1:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 7 or z[i] == 1:
            if z[1] == 2 or z[1] == 3 or z[1] == 4:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 10 or z[i] == 4:
            if z[1] == 5 or z[1] == 6 or z[1] == 7:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 6 or z[i] == 0:
            if z[1] == 8 or z[1] == 9 or z[1] == 10:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 3 or z[i] == 2:
            if z[1] == 11 or z[1] == 0 or z[1] == 1:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 0 or z[i] == 2 or z[i] == 9:
            if z[1] == 2 or z[1] == 3 or z[1] == 4:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 5 or z[i] == 11 or z[i] == 10:
            if z[1] == 5 or z[1] == 6 or z[1] == 7:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 1 or z[i] == 8:
            if z[1] == 8 or z[1] == 9 or z[1] == 10:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 6 or z[i] == 0:
            if z[1] == 11 or z[1] == 0 or z[1] == 1:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 6 or z[i] == 1 or z[i] == 10 or z[i] == 4:
            if z[2] == 2 or z[2] == 5 or z[2] == 8 or z[2] == 11:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 9 or z[i] == 3 or z[i] == 6 or z[i] == 0:
            if z[2] == 4 or z[2] == 7 or z[2] == 10 or z[2] == 1:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 11 or z[i] == 5 or z[i] == 8 or z[i] == 2:
            if z[2] == 3 or z[2] == 6 or z[2] == 9 or z[2] == 0:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 5 or z[i] == 1:
            if z[1] == 2 or z[1] == 3 or z[1] == 4:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 8 or z[i] == 4:
            if z[1] == 5 or z[1] == 6 or z[1] == 7:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 11 or z[i] == 7:
            if z[1] == 8 or z[1] == 9 or z[1] == 10:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 10 or z[i] == 2:
            if z[1] == 11 or z[1] == 0 or z[1] == 1:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")

    for i in range(4):
        if z[i] == 8:
            if z[1] == 2 or z[1] == 8:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 9:
            if z[1] == 3 or z[1] == 9:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 10:
            if z[1] == 4 or z[1] == 10:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 11:
            if z[1] == 5 or z[1] == 11:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 0:
            if z[1] == 6 or z[1] == 0:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
        if z[i] == 1:
            if z[1] == 7 or z[1] == 1:
                if i == 0:
                    listy.append("?????????")
                elif i == 1:
                    listm.append("?????????")
                elif i == 2:
                    listd.append("?????????")
                else:
                    listh.append("?????????")
    shansha[0] = shansha[0] + listy
    shansha[2] = shansha[2] + listd
    shansha[3] = shansha[3] + listh
    for i in shansha:
        i = list(set(i))


def place(table, x, y, str, isVerticle, size=15, setWid=False):
    if isVerticle:
        str = '\n'.join(str[i] for i in range(0, len(str)))
    table.cell(x, y).text = str
    table.cell(
        x, y).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    for run in table.cell(x, y).paragraphs[0].runs:
        font = run.font
        font.size = Pt(size)
    if setWid:
        table.cell(x, y).width = Inches(2)


def makeWord():
    doc = Document()  # ????????????
    table = doc.add_table(rows=12, cols=5, style='Table Grid')
    row = 0
    table.cell(row, 0).merge(table.cell(row, 4))
    table.cell(1, 0).merge(table.cell(1, 4))
    place(table, row, 0, f"{name}???????????????", False)
    row += 1
    table.cell(row, 0).merge(table.cell(row, 4))
    if(isMan):
        gender = "???"
    else:
        gender = "???"
    animal = "????????????????????????????????????"
    place(table, row, 0,
          f"??????: {gender}  ??????: {animal[yourAnimal]}  ??????: {animal[couterAnimal]}", False)
    row += 1
    table.cell(row, 0).merge(table.cell(row, 4))
    place(table, row, 0, bornDay[0], False)
    row += 1
    table.cell(row, 0).merge(table.cell(row, 4))
    place(table, row, 0, bornDay[1], False)
    row += 1
    if len(sixCan) != 0:
        table.add_row()
        table.cell(row, 0).merge(table.cell(row, 4))
        place(table, row, 0, sixCan, False)
        row += 1
    for i in range(4):
        tmp_l = ["??????", "??????", "??????", "??????"]
        place(table, row, i, tmp_l[i], True)
    row += 1
    place(table, row, 4, "??????", True, 12)
    for i in range(3, -1, -1):
        place(table, row, i, mainStar[3-i], True)
    row += 1
    place(table, row, 4, "??????", True, 12)
    for i in range(3, -1, -1):
        toColor = {"???": RGBColor(255, 55, 55),
                   "???": RGBColor(255, 227, 132),
                   "???": RGBColor(34, 139, 34),
                   "???": RGBColor(30, 144, 255),
                   "???": RGBColor(210, 180, 140),
                   }
        table.cell(row, i).text = eightword_f[3-i]
        table.cell(row, i).add_paragraph('    '+eightword_f_elem[3-i])
        table.cell(row, i).add_paragraph(eightword_s[3-i])
        table.cell(row, i).add_paragraph('    '+eightword_s_elem[3-i])
        for j in table.cell(row, i).paragraphs:
            j.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(row, i).paragraphs[0].runs[0].font.size = Pt(20)
        table.cell(row, i).paragraphs[0].runs[0].font.bold = True
        table.cell(row, i).paragraphs[1].runs[0].font.size = Pt(10)
        table.cell(row, i).paragraphs[2].runs[0].font.size = Pt(20)
        table.cell(row, i).paragraphs[2].runs[0].font.bold = True
        table.cell(row, i).paragraphs[3].runs[0].font.size = Pt(10)
    row += 1
    place(table, row, 4, "???", True, 12)
    for i in range(3, -1, -1):
        tmp = ""
        for j in range(3):
            if bury[3-i][j] != "":
                tmp += bury[3-i][j]
            else:
                tmp += '???'
            tmp += ' '
        place(table, row, i, tmp, False)
    row += 1
    place(table, row, 4, "??????", True, 12)
    for i in range(3, -1, -1):
        tmp = ""
        for j in range(2):
            for k in range(3):
                if secondStar[3-i][k] != '':
                    tmp += secondStar[3-i][k][j]
                else:
                    tmp += '???'
                tmp += ' '
            tmp += '\n'
        place(table, row, i, tmp, False)
    row += 1
    place(table, row, 4, "???", True, 12)
    for i in range(3, -1, -1):
        place(table, row, i, fate[3-i], True)
    row += 1
    place(table, row, 4, "??????", True, 12)
    for i in range(3, -1, -1):
        tmp = ""
        for j in shansha[3-i]:
            tmp += j+'\n'
        place(table, row, i, tmp, False, 12)
    row += 1
    if len(othershansha) != 0:
        table.add_row()
        table.cell(row, 0).merge(table.cell(row, 4))
        tmp = ""
        for j in othershansha:
            tmp += j+'  '
        place(table, row, 0, f"????????????: {tmp}", False)
        row += 1
    table.cell(row, 0).merge(table.cell(row, 1))
    place(table, row, 0, loveGod, False)
    table.cell(row, 2).merge(table.cell(row, 3))
    place(table, row, 2, body, False)

    table2 = doc.add_table(rows=3, cols=11, style='Table Grid')
    table2.cell(0, 10).merge(table2.cell(1, 10))
    place(table2, 0, 10, "??????", False, 12)
    bf = int(bigFate)
    for i in range(9, -1, -1):
        place(table2, 0, i, str(bf), False)
        bf += 10
    for i in range(9, -1, -1):
        place(table2, 1, i, bigFateYears[9-i], True)
    table2.cell(2, 0).merge(table2.cell(2, 10))
    place(table2, 2, 0, bigFateText, False)
    doc.add_page_break()
    table3 = doc.add_table(rows=37, cols=10, style='Table Grid')
    table3.cell(0, 0).merge(table3.cell(0, 9))
    place(table3, 0, 0, "??????", False, 10)
    row = 1
    cnt = int(bigFate)
    for i in range(1, 37, 4):
        for j in range(10):
            place(table3, i, j, str(cnt+year - 1912), False, 10)
            place(table3, i+1, j, str(cnt), False, 10)
            place(table3, i+2, j, flowYears[cnt], True, 10)
            cnt += 1
    doc.add_page_break()
    table5 = doc.add_table(rows=1, cols=2, style='Table Grid')
    row = 0
    place(table5, row, 0, "??????", False, 15, True)
    place(table5, row, 1, "??????", False, 15, True)
    row += 1
    for k, v in shanshaExplain.items():
        table5.add_row()
        place(table5, row, 0, k, False, 12, True)
        place(table5, row, 1, v, False, 12, True)
        row += 1
    doc.add_page_break()
    table4 = doc.add_table(rows=11, cols=2, style='Table Grid')
    row = 0

    place(table4, row, 0, "??????", False, 15, True)
    place(table4, row, 1, "??????", False, 15, True)
    row += 1
    for k, v in personality.items():
        place(table4, row, 0, k, False, 12, True)
        place(table4, row, 1, v, False, 12, True)
        row += 1
    table4.cell(row, 0).merge(table4.cell(row, 1))
    place(table4, row, 0, " ", False, 15, True)
    row += 1
    place(table4, row, 0, "?????????", False, 15, True)
    place(table4, row, 1, "??????", False, 15, True)
    row += 1
    for k, v in horseFlower.items():
        place(table4, row, 0, k, False, 12, True)
        place(table4, row, 1, v, False, 12, True)
        row += 1
    
    doc.save(f'{name}.docx')
    pythoncom.CoInitialize()
    convert(f'{name}.docx', f'./data/{name}.pdf')



def send(name):
    # coding=utf-8

    file = '.\\data\\'+str(name)+'.pdf'
    server_address = 'smtp.gmail.com'
    port = 587
    send_user = 'why33551@gmail.com'
    receive_users = 'mama33551@gmail.com'
    subject = f'{name}???????????????'
    password = "tyevkckncblwgnut"

    mail_type = '1'

    msg = MIMEMultipart()
    msg['Subject'] = subject
    msg['From'] = send_user
    msg['To'] = receive_users

    part_text = MIMEText('?????????????????????')
    msg.attach(part_text)

    part_attach1 = MIMEApplication(open(file, 'rb').read())
    part_attach1.add_header('Content-Disposition',
                            'attachment', filename=file)
    msg.attach(part_attach1)

    smtp = smtplib.SMTP(server_address, port)
    smtp.ehlo()  # ??????SMTP?????????
    smtp.starttls()
    smtp.login(send_user, password)
    smtp.sendmail(send_user, receive_users, msg.as_string())
    print('?????????????????????')
    


def sendGUI():
    global name
    global visited
    global win
    win.destroy()
    win = tk.Tk()
    setupWin()
    name_l = tk.Label(text="??????", font="??????????????? 20")
    name_c = ttk.Combobox(win, font="??????????????? 20", state="readonly")
    recnames = []
    for file in glob.glob(".\\Words\\*.docx"):
        recnames.append(str(file)[8:-5])
    name_c['value'] = recnames
    name_c.current(0)
    name_b = tk.Button(
        text="??????", command=lambda: send(name_c.get()), font="??????????????? 15")
    name_l.grid(row=0, column=0)
    name_c.grid(row=0, column=1)
    name_b.grid(row=0, column=2)


def dele(name):
    win.destroy()
    file = '.\\Words\\'+str(name)+'.docx'
    os.remove(file)
    quit()


def delGUI():
    global name
    global visited
    global win
    win.destroy()
    win = tk.Tk()
    setupWin()
    name_l = tk.Label(text="??????", font="??????????????? 20")
    name_c = ttk.Combobox(win, font="??????????????? 20", state="readonly")
    recnames = []
    for file in glob.glob(".\\Words\\*.docx"):
        recnames.append(str(file)[8:-5])
    name_c['value'] = recnames
    name_c.current(0)
    name_b = tk.Button(
        text="??????", command=lambda: dele(name_c.get()), font="??????????????? 15")
    name_l.grid(row=0, column=0)
    name_c.grid(row=0, column=1)
    name_b.grid(row=0, column=2)


def menuGUI():
    setupWin()
    add_b = tk.Button(
        text="??????", command=GUI, font="??????????????? 15")
    add_b.pack()
    send_b = tk.Button(
        text="??????E-mail", command=sendGUI, font="??????????????? 15")
    send_b.pack()
    del_b = tk.Button(
        text="??????", command=delGUI, font="??????????????? 15", bg="red")
    del_b.pack()
    win.mainloop()


# main

name = " "
year = 0
mon = 0
day = 0
hr = 0

isMan = True
isLunar = False
isLeapMonth = False



mainStar = []
bornDay = []
eightword_f = []
eightword_s = []
eightword_f_elem = []
eightword_s_elem = []
bury = [[], [], [], []]
secondStar = [[], [], [], []]
fate = []
shansha = [[], [], [], []]
othershansha = []
bigFate = ""
bigFateYears = []
bigFateText = ""
flowYears = []
flowYears.append(" ")  # index 1
loveGod = ""
body = ""
personality = {}
horseFlower = {}
shanshaExplain = {}
couterAnimal = 0
yourAnimal = 0
sixCan = ""
def generate(nam, yea, mo, da, h, isMa, isLuna, isLeapMont):
    global name
    global year
    global mon
    global day
    global hr
    global isMan
    global isLunar
    global isLeapMonth
    global mainStar 
    global bornDay 
    global eightword_f 
    global eightword_s 
    global eightword_f_elem 
    global eightword_s_elem 
    global bury 
    global secondStar
    global fate 
    global shansha
    global othershansha 
    global bigFate 
    global bigFateYears 
    global bigFateText
    global flowYears 
    global loveGod
    global body 
    global personality
    global horseFlower 
    global shanshaExplain 
    global couterAnimal 
    global yourAnimal 
    global sixCan
    mainStar = []
    bornDay = []
    eightword_f = []
    eightword_s = []
    eightword_f_elem = []
    eightword_s_elem = []
    bury = [[], [], [], []]
    secondStar = [[], [], [], []]
    fate = []
    shansha = [[], [], [], []]
    othershansha = []
    bigFate = ""
    bigFateYears = []
    bigFateText = ""
    flowYears = []
    flowYears.append(" ")  # index 1
    loveGod = ""
    body = ""
    personality = {}
    horseFlower = {}
    shanshaExplain = {}
    couterAnimal = 0
    yourAnimal = 0
    sixCan = ""
    name = nam
    year = yea
    mon = mo
    day = da
    hr = h

    isMan = isMa
    isLunar = isLuna
    isLeapMonth = isLeapMont

    crawler()
    adjust()

    makeWord()
